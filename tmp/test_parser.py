import sys
import os

# Add project root to path
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from bot.services.parser import DocxParser
from docx import Document

# Create a dummy docx
doc = Document()
doc.add_paragraph("? Question 1")
doc.add_paragraph("+ Correct A")
doc.add_paragraph("= Correct B (Equal)")
doc.add_paragraph("- Incorrect C")

test_file = "test_markers.docx"
doc.save(test_file)

parser = DocxParser(test_file)
blocks, errors = parser.parse()

print(f"Blocks found: {len(blocks)}")
if blocks:
    q = blocks[0]
    print(f"Question: {''.join(p.text for p in q.question_paragraphs)}")
    for i, ans in enumerate(q.answers):
        text = "".join(p.text for p in ans.original_paragraphs)
        print(f"Ans {i+1}: {text} (is_correct={ans.is_correct})")

os.remove(test_file)
