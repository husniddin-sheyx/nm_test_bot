import sys
import os

# Add project root to path
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from bot.services.parser import DocxParser
from bot.services.validator import Validator
from docx import Document

# Create a dummy docx with multiple correct answers and numbered options
doc = Document()
doc.add_paragraph("? Question Medicine")
doc.add_paragraph("1) Option One")
doc.add_paragraph("2) Option Two")
doc.add_paragraph("3) Option Three")
doc.add_paragraph("+ 1,2")
doc.add_paragraph("= 2,3")
doc.add_paragraph("- 4 (Incorrect)")

test_file = "test_medicine.docx"
doc.save(test_file)

parser = DocxParser(test_file)
blocks, errors = parser.parse()

validator = Validator()
valid, invalid = validator.validate(blocks)

print(f"Blocks parsed: {len(blocks)}")
print(f"Valid: {len(valid)}")
print(f"Invalid errors: {len(invalid)}")

if valid:
    q = valid[0]
    print(f"Question text count: {len(q.question_paragraphs)}")
    for p in q.question_paragraphs:
        print(f"  Q: {p.text}")
    print(f"Answers count: {len(q.answers)}")
    for i, a in enumerate(q.answers):
        print(f"  A{i+1}: {a.original_paragraphs[0].text} (Correct={a.is_correct})")

os.remove(test_file)
